"""
文档转换模块 - 将DOCX/DOC/PDF转换为Markdown
"""

import os
from typing import Optional

# 检查依赖库
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from poword.api.word import doc2docx
    POWORD_AVAILABLE = True
except ImportError:
    POWORD_AVAILABLE = False
    doc2docx = None


class DocumentConverter:
    """文档转换器"""
    
    def convert(self, file_path: str) -> Optional[str]:
        """自动识别并转换文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            转换后的Markdown内容
        """
        if not os.path.exists(file_path):
            print(f"    [X] 文件不存在: {file_path}")
            return None
        
        # 根据扩展名选择转换方法
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.docx':
            return self.docx_to_markdown(file_path)
        elif ext == '.doc':
            return self.doc_to_markdown(file_path)
        elif ext == '.pdf':
            return self.pdf_to_markdown(file_path)
        else:
            print(f"    [X] 不支持的文件格式: {ext}")
            return None
    
    def docx_to_markdown(self, docx_path: str) -> Optional[str]:
        """将DOCX文件转换为Markdown
        
        Args:
            docx_path: DOCX文件路径
            
        Returns:
            Markdown内容
        """
        if not DOCX_AVAILABLE:
            print("    [X] python-docx未安装，无法转换DOCX")
            return None
        
        try:
            doc = Document(docx_path)
            markdown_lines = []
            
            # 处理段落
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    markdown_lines.append('')
                    continue
                
                # 根据样式判断标题级别
                style_name = paragraph.style.name if paragraph.style else ''
                
                if 'Heading 1' in style_name or '标题 1' in style_name:
                    markdown_lines.append(f'# {text}')
                elif 'Heading 2' in style_name or '标题 2' in style_name:
                    markdown_lines.append(f'## {text}')
                elif 'Heading 3' in style_name or '标题 3' in style_name:
                    markdown_lines.append(f'### {text}')
                elif 'Heading 4' in style_name or '标题 4' in style_name:
                    markdown_lines.append(f'#### {text}')
                elif 'Heading 5' in style_name or '标题 5' in style_name:
                    markdown_lines.append(f'##### {text}')
                elif 'Heading 6' in style_name or '标题 6' in style_name:
                    markdown_lines.append(f'###### {text}')
                else:
                    # 处理段落中的格式
                    para_text = self._process_paragraph_runs(paragraph)
                    if para_text:
                        markdown_lines.append(para_text)
            
            # 处理表格
            for table in doc.tables:
                markdown_lines.append('')
                markdown_lines.append(self._table_to_markdown(table))
                markdown_lines.append('')
            
            content = '\n'.join(markdown_lines)
            print(f"    [OK] DOCX转换成功，内容长度: {len(content)} 字符")
            return content
            
        except Exception as e:
            print(f"    [X] DOCX转换失败: {e}")
            return None
    
    def _process_paragraph_runs(self, paragraph) -> str:
        """处理段落中的run格式"""
        text_parts = []
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # 处理格式
            if run.bold:
                text = f'**{text}**'
            if run.italic:
                text = f'*{text}*'
            if run.underline:
                text = f'<u>{text}</u>'
            
            text_parts.append(text)
        
        return ''.join(text_parts)
    
    def _table_to_markdown(self, table) -> str:
        """将表格转换为Markdown格式"""
        markdown_lines = []
        
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            markdown_lines.append('| ' + ' | '.join(cells) + ' |')
            
            # 添加表头分隔线
            if i == 0:
                markdown_lines.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
        
        return '\n'.join(markdown_lines)
    
    def pdf_to_markdown(self, pdf_path: str) -> Optional[str]:
        """将PDF文件转换为Markdown
        
        Args:
            pdf_path: PDF文件路径
            
        Returns:
            Markdown内容
        """
        if not PDF_AVAILABLE:
            print("    [X] pypdf未安装，无法提取PDF文本")
            return None
        
        try:
            reader = PdfReader(pdf_path)
            markdown_lines = []
            
            total_pages = len(reader.pages)
            print(f"    PDF页数: {total_pages}")
            
            extracted_text_count = 0
            
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        extracted_text_count += 1
                        lines = text.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line:
                                markdown_lines.append(line)
                        markdown_lines.append('')
                except Exception as e:
                    print(f"    页面 {page_num} 提取失败: {e}")
                    continue
            
            if extracted_text_count == 0:
                print("    [X] PDF可能是扫描版，无法提取文本（需要OCR）")
                return None
            
            content = '\n'.join(markdown_lines).strip()
            if len(content) > 100:
                print(f"    [OK] 成功提取 {extracted_text_count}/{total_pages} 页文本")
                return content
            else:
                print(f"    [X] 提取的文本内容过少: {len(content)} 字符")
                return None
            
        except Exception as e:
            print(f"    [X] PDF提取失败: {e}")
            return None
    
    def doc_to_markdown(self, doc_path: str) -> Optional[str]:
        """将DOC文件转换为Markdown（使用poword库）
        
        Args:
            doc_path: DOC文件路径
            
        Returns:
            Markdown内容
        """
        if not os.path.exists(doc_path):
            print(f"    [X] 文件不存在: {doc_path}")
            return None
        
        # 使用poword转换DOC -> DOCX -> Markdown
        if not POWORD_AVAILABLE:
            print("    [X] poword未安装，无法转换DOC文件")
            print("    请安装: pip install poword")
            return None
        
        try:
            import tempfile
            from poword.api.word import doc2docx
            
            with tempfile.TemporaryDirectory() as tmpdir:
                output_name = 'converted.docx'
                docx_path = os.path.join(tmpdir, output_name)
                
                # 使用poword将DOC转换为DOCX
                doc2docx(doc_path, tmpdir, output_name)
                
                if os.path.exists(docx_path):
                    # 将DOCX转换为Markdown
                    content = self.docx_to_markdown(docx_path)
                    if content:
                        print("    [OK] 使用poword转换DOC成功")
                        return content
                    else:
                        print("    [X] DOCX转换失败")
                        return None
                else:
                    print("    [X] poword转换失败，未生成DOCX文件")
                    return None
                    
        except Exception as e:
            print(f"    [X] DOC转换失败: {e}")
            return None

